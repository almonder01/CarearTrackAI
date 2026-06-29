from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "CareerTrackAI_Project_Report_AR.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 30, 45)
MUTED = RGBColor(90, 100, 115)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "EAF7F4"


def set_bidi_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths):
            set_cell_width(row.cells[idx], width)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def style_run(run, bold=False, italic=False, color=None, size=None):
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)


def add_paragraph(doc, text="", style=None, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p = doc.add_paragraph(style=style)
    set_bidi_paragraph(p, align)
    if text:
        r = p.add_run(text)
        style_run(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_bidi_paragraph(p)
    r = p.add_run(text)
    style_run(r, bold=True, color=BLUE if level < 3 else DARK_BLUE, size={1: 16, 2: 13, 3: 12}.get(level, 12))
    return p


def add_label_paragraph(doc, label, text):
    p = add_paragraph(doc)
    r1 = p.add_run(label)
    style_run(r1, bold=True, color=DARK_BLUE)
    r2 = p.add_run(" " + text)
    style_run(r2)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_bidi_paragraph(p)
    r = p.add_run(text)
    style_run(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_bidi_paragraph(p)
    r = p.add_run(text)
    style_run(r)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    set_bidi_paragraph(p)
    rt = p.add_run(title)
    style_run(rt, bold=True, color=DARK_BLUE)
    rb = p.add_run("\n" + body)
    style_run(rb)
    add_paragraph(doc, "")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        set_cell_shading(hdr[idx], LIGHT_FILL)
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[idx].paragraphs[0]
        set_bidi_paragraph(p)
        r = p.add_run(text)
        style_run(r, bold=True, color=DARK_BLUE)

    for row_values in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            row_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = row_cells[idx].paragraphs[0]
            set_bidi_paragraph(p)
            r = p.add_run(str(value))
            style_run(r, size=10.5)
    add_paragraph(doc, "")
    return table


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("CareerTrackAI Project Report")
    style_run(r, color=MUTED, size=9)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, color in [(1, 16, BLUE), (2, 13, BLUE), (3, 12, DARK_BLUE)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt({1: 16, 2: 12, 3: 8}[level])
        style.paragraph_format.space_after = Pt({1: 8, 2: 6, 3: 4}[level])
        style.paragraph_format.line_spacing = 1.10


def build_document():
    doc = Document()
    configure_document(doc)

    # Cover page
    p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("تقرير مشروع CareerTrackAI")
    style_run(r, bold=True, color=BLUE, size=24)
    p.paragraph_format.space_after = Pt(12)

    p = add_paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("منصة ذكية لإدارة البحث عن التدريب والوظائف باستخدام الذكاء الاصطناعي")
    style_run(r, color=INK, size=14)

    add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(
        doc,
        "ملخص سريع",
        "CareerTrackAI هو نظام ويب متكامل يساعد الطالب أو الباحث عن عمل على اكتشاف الفرص، حفظ الشركات، تتبع حالة التقديم، إدارة السير الذاتية، التحضير للمقابلات، واستخدام الذكاء الاصطناعي كدليل عملي داخل سير العمل.",
    )
    add_label_paragraph(doc, "نوع المشروع:", "تطبيق ويب Full-Stack باستخدام ASP.NET Core و React.")
    add_label_paragraph(doc, "لغة واجهة المستخدم:", "الإنجليزية.")
    add_label_paragraph(doc, "نطاق التقرير:", "شرح المشروع، الفئة المستهدفة، المعمارية، الأكواد، واجهات API، التحديات والحلول، والتحسينات المستقبلية.")
    add_label_paragraph(doc, "تاريخ التقرير:", "يونيو 2026.")
    doc.add_page_break()

    add_heading(doc, "فهرس المحتويات", 1)
    for item in [
        "1. المقدمة",
        "2. فكرة المشروع والمشكلة التي يحلها",
        "3. الفئة المستهدفة وأهداف النظام",
        "4. المعمارية العامة وسير العمل",
        "5. التقنيات والأدوات المستخدمة",
        "6. شرح مختصر للباك اند",
        "7. شرح مختصر للفرونت اند",
        "8. الذكاء الاصطناعي وواجهات API الخارجية",
        "9. أهم التحديات والحلول",
        "10. الاختبارات والتحقق",
        "11. الإضافات المستقبلية",
        "12. الخاتمة",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_heading(doc, "1. المقدمة", 1)
    add_paragraph(
        doc,
        "يمر الطالب أو الباحث عن تدريب بعدد كبير من الخطوات المتفرقة: البحث عن الشركات، حفظ روابط الوظائف، تجهيز السيرة الذاتية، متابعة حالة كل طلب، معرفة مواعيد المقابلات، والرجوع إلى الملاحظات عند الحاجة. غالبًا تتم هذه العملية عبر ملفات Excel أو ملاحظات شخصية أو روابط مبعثرة، مما يؤدي إلى فقدان الفرص أو تكرار التقديم أو ضعف المتابعة.",
    )
    add_paragraph(
        doc,
        "من هنا تأتي فكرة CareerTrackAI كنظام موحد يجمع هذه العملية في مساحة عمل واحدة، ويضيف طبقة ذكاء اصطناعي تساعد المستخدم في التحليل، الترتيب، التلخيص، والتحضير. المشروع لا يكتفي بعرض معلومات ثابتة، بل يحاول تحويل البحث عن الوظيفة إلى عملية منظمة قابلة للقياس والتحسين.",
    )

    add_heading(doc, "2. فكرة المشروع والمشكلة التي يحلها", 1)
    add_paragraph(
        doc,
        "المشروع عبارة عن منصة Web Application لإدارة رحلة البحث عن التدريب أو الوظيفة من البداية إلى النهاية. يستطيع المستخدم إدخال أو استيراد الشركات والفرص، مراجعة البيانات قبل حفظها، تحويل الفرصة إلى Application، نقلها بين مراحل مختلفة مثل Planning وApplied وInterview وAccepted وRejected، ثم إدارة المقابلات والسير الذاتية والتحليلات المرتبطة بها.",
    )
    add_table(
        doc,
        ["المشكلة", "حل CareerTrackAI"],
        [
            ["تشتت بيانات الشركات والوظائف", "تجميع الشركات والفرص في Data Hub و Opportunities مع دعم CSV و APIs خارجية."],
            ["صعوبة متابعة حالات التقديم", "Kanban Board للتطبيقات يوضح كل مرحلة ويمكن نقل الطلب بينها."],
            ["ضعف الاستفادة من السيرة الذاتية", "رفع CV وتحليلها بالذكاء الاصطناعي واستخراج نقاط القوة والضعف."],
            ["نسيان المقابلات أو نقص التحضير", "صفحة Interviews لحفظ الموعد والرابط والملاحظات وتوليد AI prep notes."],
            ["عدم وضوح أفضل فرصة للتقديم", "لوحات AI مدمجة تعطي توصيات عند الطلب وتحفظ آخر نتيجة."],
        ],
        [3000, 6360],
    )

    add_heading(doc, "3. الفئة المستهدفة وأهداف النظام", 1)
    add_heading(doc, "الفئة المستهدفة", 2)
    for text in [
        "طلاب الجامعات الباحثون عن تدريب أو وظيفة أولى.",
        "الخريجون الجدد الذين يحتاجون إلى تنظيم طلباتهم ومتابعة المقابلات.",
        "مراكز الإرشاد المهني في الجامعات التي ترغب في مساعدة الطلاب على إدارة رحلتهم المهنية.",
        "أي مستخدم يريد بناء قاعدة شركات وفرص خاصة به بدل الاعتماد على بيانات عامة فقط.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "الأهداف الرئيسية", 2)
    for text in [
        "توفير مساحة واحدة لإدارة الشركات والفرص والتقديمات والسير الذاتية والمقابلات.",
        "دمج الذكاء الاصطناعي في نقاط عملية: تحليل السيرة، توصيات، تحضير مقابلة، فحص روابط، ومساعدة المستخدم داخل النظام.",
        "دعم مصادر خارجية مثل Adzuna و JobDataLake مع مراجعة البيانات قبل الاستيراد.",
        "حماية خصوصية المستخدم بجعل البيانات الخاصة مرتبطة بحسابه، مع إمكانية حفظ بيانات مشتركة عند الحاجة فقط.",
        "تقديم واجهة حديثة تدعم Dark Mode وتخزين تفضيلات المستخدم.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "4. المعمارية العامة وسير العمل", 1)
    add_paragraph(
        doc,
        "يعتمد المشروع على معمارية Client-Server. الفرونت اند مبني بـ React + Vite ويتصل بالباك اند عبر Axios. الباك اند مبني بـ ASP.NET Core Web API ويتعامل مع SQL Server عبر Entity Framework Core. كما توجد خدمات منفصلة للتعامل مع Gemini وAdzuna وJobDataLake واستيراد/تصدير البيانات.",
    )
    add_table(
        doc,
        ["الطبقة", "المسؤولية"],
        [
            ["Frontend", "عرض الواجهة، إدارة الحالة، التنقل، النماذج، Dark Mode، واستدعاء API."],
            ["Backend API", "تنفيذ منطق الأعمال، المصادقة، التحكم في البيانات، وتقديم endpoints."],
            ["Database", "حفظ المستخدمين والشركات والفرص والتقديمات والسير والمقابلات والتنبيهات."],
            ["AI & External APIs", "توفير التحليل والتوصيات واستيراد الفرص من مصادر خارجية."],
        ],
        [2200, 7160],
    )
    add_heading(doc, "سيناريو استخدام مختصر", 2)
    for text in [
        "يسجل المستخدم حسابًا جديدًا ويكمل بياناته الشخصية.",
        "يستخدم Data Hub لجلب أو رفع بيانات شركات وفرص، ثم يراجع الصفوف قبل الاستيراد.",
        "تظهر الفرص المحفوظة في Opportunities، ويمكنه فتح الرابط أو تتبعها كتقديم.",
        "ينتقل الطلب إلى Applications حيث يغير الحالة حسب تقدمه الفعلي.",
        "إذا وصل الطلب إلى Interview، تظهر بطاقة في Interviews ليضيف الموعد والرابط والملاحظات.",
        "يرفع CV في Resumes ويحللها بالذكاء الاصطناعي ثم يستخدم AI Studio للمساعدة العامة.",
    ]:
        add_number(doc, text)

    add_heading(doc, "5. التقنيات والأدوات المستخدمة", 1)
    add_table(
        doc,
        ["الجزء", "التقنيات"],
        [
            ["Backend", "ASP.NET Core Web API, .NET 10, Entity Framework Core, SQL Server, JWT Authentication, BCrypt, Scalar/OpenAPI."],
            ["Frontend", "React 19, Vite 7, Tailwind CSS 4, React Router, Axios, lucide-react, Recharts, dayjs, react-markdown."],
            ["AI", "Google Gemini API عبر Google AI Studio، مع طبقة fallback محلية عند عدم توفر المفتاح."],
            ["External Jobs APIs", "Adzuna API و JobDataLake API، إضافة إلى AI sourcing عبر Google/LinkedIn scouting logic."],
            ["Development", "npm scripts, oxlint, EF Core migrations, dotnet user-secrets, Vite dev server."],
        ],
        [2200, 7160],
    )

    add_heading(doc, "6. شرح مختصر للباك اند", 1)
    add_paragraph(
        doc,
        "الباك اند يمثل قلب النظام؛ فهو المسؤول عن قواعد البيانات، المصادقة، منطق الأعمال، التكامل مع الذكاء الاصطناعي، واستيراد البيانات من المصادر الخارجية. تم تقسيمه إلى Controllers وServices وDTOs وModels وData Context، مما يجعل الكود أوضح وأسهل في الصيانة.",
    )
    add_heading(doc, "أهم المكونات", 2)
    add_table(
        doc,
        ["المكون", "الدور"],
        [
            ["AppDbContext", "تعريف DbSet للجداول مثل Users, Companies, JobOpportunities, Applications, Resumes, Interviews, Notifications."],
            ["AuthController / AuthService", "تسجيل الدخول، إنشاء الحساب، JWT، refresh tokens، وتشفير كلمات المرور باستخدام BCrypt."],
            ["CompaniesController / CompanyService", "إدارة الشركات، الاستيراد، التصدير، الحفظ من قاعدة مشتركة إلى مساحة المستخدم."],
            ["JobOpportunitiesController / JobOpportunityService", "إدارة الفرص، الفلاتر، الاستيراد من CSV، Adzuna، JobDataLake، وAI sourcing."],
            ["ApplicationsController / ApplicationService", "تتبع حالة التقديم وتغييرها وحذف الطلب من Pipeline."],
            ["ResumesController / ResumeService", "رفع السير الذاتية وحفظ ملفات PDF/DOCX وربطها بالمستخدم."],
            ["ResumeTextExtractionService", "استخراج النص من DOCX وPDF قدر الإمكان لاستخدامه في تحليل السيرة."],
            ["InterviewService", "إنشاء وتعديل وحذف المقابلات وحفظ التاريخ والمدة والرابط والملاحظات."],
            ["AiService", "المحادثة، تحليل السيرة، توليد خطاب التقديم، التوصيات، فحص Gemini token."],
            ["UsageController", "عرض استخدام Gemini وواجهات APIs الخارجية."],
        ],
        [3000, 6360],
    )
    add_heading(doc, "نقاط مهمة في منطق الباك اند", 2)
    for text in [
        "تم ربط البيانات الحساسة بالمستخدم من خلال UserId حتى لا يرى حساب جديد بيانات حساب آخر.",
        "تم دعم Shared Database بشكل اختياري؛ لا تدخل بياناتها إلى مساحة المستخدم إلا عند الضغط على حفظ.",
        "تم استخدام User Secrets لتخزين مفتاح Gemini أثناء التطوير بدل وضعه مباشرة في الكود.",
        "تم تحسين رسائل اختبار Gemini token حتى لا تظهر رسائل تقنية غير مفهومة للمستخدم.",
        "تم توفير endpoints للمقابلات تسمح بالإضافة والتعديل والحذف دون الحاجة إلى تغيير قاعدة البيانات.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "7. شرح مختصر للفرونت اند", 1)
    add_paragraph(
        doc,
        "الفرونت اند يمثل تجربة المستخدم الكاملة، وتم بناؤه باللغة الإنجليزية كما هو مطلوب. تم الاعتماد على React Router للتنقل، Axios للتواصل مع الباك اند، Tailwind CSS للتصميم، وlocalStorage لحفظ بعض التفضيلات والنتائج المؤقتة مثل إعدادات الواجهة ونتائج AI panels.",
    )
    add_table(
        doc,
        ["الصفحة", "الوظيفة"],
        [
            ["Dashboard", "ملخص عام، مؤشرات الأداء، Getting Started، ورسوم بيانية يمكن إخفاؤها من الإعدادات."],
            ["Applications", "Kanban board لتتبع الطلبات بين Planning, Applied, Interview, Accepted, Rejected مع حذف الطلب."],
            ["Opportunities", "عرض الفرص، فلترة المصادر، فتح رابط الوظيفة، تتبع الفرصة، فحص الرابط بالذكاء الاصطناعي."],
            ["Resumes", "رفع CV، تحليلها بالذكاء الاصطناعي، حفظ نتيجة التحليل، وحذف السير غير المطلوبة."],
            ["Interviews", "عرض طلبات Interview، جدولة موعد، حفظ رابط الاجتماع، ملاحظات التحضير، AI prep notes."],
            ["AI Studio", "محادثة موجهة للمساعدة المهنية، توصيات، اختبار Gemini token، وتنسيق Markdown."],
            ["Data Hub", "استيراد/تصدير CSV، مراجعة preview rows، Adzuna، JobDataLake، AI sourcing، والشركات المشتركة."],
            ["Settings", "Dark/Light/System mode، تفضيلات الواجهة، إعدادات AI provider، والخطط المستقبلية."],
            ["Usage", "عرض استهلاك Gemini وAdzuna وJobDataLake مع ملخص عام."],
        ],
        [2500, 6860],
    )
    add_heading(doc, "تحسينات تجربة المستخدم", 2)
    for text in [
        "دعم Dark Mode مع تحسين ألوان النصوص والبطاقات لتكون مريحة.",
        "Sidebar قابل للطي والتمرير حتى لا تختفي آخر التبويبات.",
        "AI Copilot قابل للإغلاق والتحكم من Settings.",
        "حفظ نتائج تحليل السيرة وAI Studio وAI panels حتى لا تختفي عند التنقل.",
        "إظهار حالات loading داخل الأزرار لتوضيح أن النظام يعمل على جلب أو تحليل البيانات.",
        "استخدام react-markdown لعرض نصوص الذكاء الاصطناعي بتنسيق صحيح بدل ظهور نجوم Markdown.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "8. الذكاء الاصطناعي وواجهات API الخارجية", 1)
    add_paragraph(
        doc,
        "تم دمج الذكاء الاصطناعي ليس فقط كشعار أو نص ثابت، بل كجزء عملي داخل التجربة. توجد وظائف AI في تحليل السير الذاتية، المحادثة، توصيات Dashboard، اقتراحات Pipeline، تحضير المقابلات، فحص روابط الوظائف، والمساعدة في Data Hub.",
    )
    add_table(
        doc,
        ["API / Service", "الاستخدام داخل المشروع"],
        [
            ["Google Gemini", "المحادثة، التوصيات، تحليل CV، توليد cover letter، interview prep notes، فحص الروابط، واختبار الاتصال."],
            ["Adzuna API", "استيراد فرص عمل من الدول المدعومة، خصوصًا عند البحث عن internships أو jobs."],
            ["JobDataLake API", "جلب فرص وظائف enriched data مع دعم provider مستقل في Data Hub وUsage."],
            ["AI Sourcing Agent", "يبني خطة بحث ويستخدم provider مختار مثل JobDataLake أو Adzuna أو scouting عبر Google/LinkedIn."],
            ["Internal CareerTrackAI API", "كل عمليات النظام مثل users, companies, opportunities, applications, resumes, interviews, notifications, usage."],
        ],
        [2500, 6860],
    )
    add_heading(doc, "سياسة عرض الذكاء الاصطناعي في الواجهة", 2)
    for text in [
        "لوحات AI لا تولد النص تلقائيًا عند فتح الصفحة؛ يظهر العنوان فقط حتى يضغط المستخدم Refresh.",
        "عند الضغط على Refresh يتم استدعاء AI وتخزين النتيجة في localStorage.",
        "تظل النتيجة موجودة عند الانتقال بين الصفحات حتى يمسحها المستخدم.",
        "كل لوحة لها زر حذف يمسح نتيجة تلك اللوحة فقط.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "9. أهم التحديات والحلول", 1)
    add_table(
        doc,
        ["التحدي", "الحل الذي تم تطبيقه"],
        [
            ["ظهور بيانات حساب قديم في حساب جديد", "إضافة UserId للشركات والفرص وربط البيانات الخاصة بالمستخدم، مع جعل shared database اختيارية."],
            ["Adzuna لا يدعم بعض الدول مثل Malaysia مباشرة", "إضافة قائمة دول مدعومة واستخدام Singapore كخيار عملي قريب، مع رسائل توضيحية عند عدم وجود نتائج."],
            ["JobDataLake أعاد 500 لبعض الاستعلامات", "تحسين التعامل مع الرسائل وتوفير provider مستقل وpreview قابل للمراجعة قبل الاستيراد."],
            ["تحليل PDF للسيرة يعطي نتيجة فارغة", "إضافة ResumeTextExtractionService ومحاولة استخراج النص من DOCX/PDF، ثم fallback إلى إرسال الملف للذكاء الاصطناعي عند الحاجة."],
            ["رسائل Gemini token غير مفهومة", "تحويل الأخطاء التقنية إلى رسائل واضحة مثل quota reached أو invalid API key."],
            ["Dark Mode غير واضح في بعض الصفحات", "توحيد ألوان النصوص والبطاقات والحقول لتناسب الوضع الداكن."],
            ["نتائج AI تختفي عند التنقل", "تخزين نتائج AI Studio وResume Analysis وAI panels في localStorage حتى يمسحها المستخدم."],
            ["AI markdown يظهر كنجوم ورموز", "استخدام react-markdown وremark-gfm لعرض التنسيق بشكل صحيح."],
            ["Sidebar يخفي آخر التبويبات", "تحويله إلى تخطيط flex مع منطقة navigation قابلة للتمرير."],
            ["المقابلات كانت تعرض فقط", "إضافة نموذج جدولة وتعديل وحذف وربط AI prep notes وحفظ الرابط والملاحظات."],
        ],
        [2900, 6460],
    )

    add_heading(doc, "10. الاختبارات والتحقق", 1)
    add_paragraph(
        doc,
        "تم تنفيذ اختبارات تحقق عملية أثناء التطوير للتأكد من أن الواجهة تعمل وأن الباك اند يقبل التعديلات. أهم عمليات التحقق شملت بناء الفرونت اند، lint، بناء الباك اند إلى مجلد مؤقت عند وجود السيرفر شغال، وتجربة سيناريوهات المستخدم مثل تسجيل حساب جديد، استيراد الفرص، تتبع الطلبات، تحليل السيرة، وجدولة المقابلات.",
    )
    add_table(
        doc,
        ["نوع التحقق", "الأمر أو الطريقة", "النتيجة"],
        [
            ["Frontend lint", "npm run lint", "نجح."],
            ["Frontend build", "npm run build", "نجح مع تحذير طبيعي عن حجم bundle."],
            ["Backend build", "dotnet build إلى مجلد tmp عند قفل ملف التشغيل", "نجح عند البناء إلى مجلد مؤقت."],
            ["Database migrations", "EF Core migrations", "تمت إضافة SourceProvider وUser Scope للبيانات."],
            ["User workflow", "تجربة المسار من Opportunities إلى Applications إلى Interviews", "تم تحسين المنطق وإضافة الحذف والتعديل."],
        ],
        [2300, 4300, 2760],
    )

    add_heading(doc, "11. القيود الحالية", 1)
    for text in [
        "بوابة الدفع والخطط موجودة كتصميم وتجربة أولية، لكنها لم تربط بعد ببوابة دفع حقيقية مثل Stripe أو Moyasar أو Tap Payments.",
        "حفظ Personal Gemini API key في الواجهة ما زال مؤقتًا، والأفضل مستقبلاً تخزينه في Backend encrypted vault.",
        "استخدام Gemini يخضع لحدود quota اليومية الخاصة بالمفتاح المستخدم.",
        "بعض مزودي الوظائف لا يدعمون كل الدول أو قد يرجعون نتائج قليلة حسب الاستعلام.",
        "حجم bundle في Vite كبير نسبيًا، ويمكن تحسينه لاحقًا عبر code splitting.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "12. الإضافات المستقبلية المقترحة", 1)
    add_table(
        doc,
        ["الإضافة", "الفائدة المتوقعة"],
        [
            ["بوابة دفع حقيقية", "تفعيل الخطط المدفوعة وإدارة الاشتراكات بشكل كامل."],
            ["Backend vault لمفاتيح API", "حفظ مفاتيح المستخدمين بشكل آمن ومشفر بدل localStorage."],
            ["Email / Calendar reminders", "إرسال تنبيهات للمواعيد والمقابلات وال deadlines."],
            ["Admin dashboard", "إدارة المستخدمين، المصادر المشتركة، ومراقبة استخدام APIs."],
            ["RAG / Vector Search", "استخدام Pinecone أو بديل مشابه لجعل AI يفهم ملفات المستخدم وسياق المشروع بعمق أكبر."],
            ["Deployment pipeline", "نشر المشروع على Azure أو VPS مع CI/CD وبيئة Production واضحة."],
            ["Advanced analytics", "تحليل معدل النجاح، أفضل مصادر الفرص، وأسباب الرفض أو القبول."],
            ["Mobile responsive polish", "تحسينات إضافية لشاشات الهاتف وتجربة المستخدم السريعة."],
        ],
        [3000, 6360],
    )

    add_heading(doc, "13. الخاتمة", 1)
    add_paragraph(
        doc,
        "يمثل CareerTrackAI مشروعًا عمليًا متكاملًا يجمع بين إدارة البيانات، تجربة المستخدم الحديثة، والذكاء الاصطناعي الموجه لحل مشكلة حقيقية. القوة الأساسية في المشروع ليست في وجود صفحات متعددة فقط، بل في ربطها بسير عمل واضح: البحث عن الفرص، مراجعتها، حفظها، تتبعها، تجهيز السيرة، التحضير للمقابلة، ثم قياس الاستخدام والتحسين.",
    )
    add_paragraph(
        doc,
        "المشروع جاهز للاستخدام في نطاقه الحالي ما عدا جزء الدفع والاشتراكات الذي يمكن إضافته لاحقًا. كما أن البنية الحالية تسمح بتوسعات مستقبلية مهمة مثل إدارة المفاتيح بشكل آمن، ربط بوابات الدفع، إضافة تذكيرات بريدية، وبناء لوحة إدارة. وبذلك يشكل المشروع أساسًا قويًا لمنصة Career Operations ذكية وقابلة للتطوير.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_document()
    print(OUT)
