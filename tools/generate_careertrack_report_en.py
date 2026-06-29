from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "CareerTrackAI_Project_Report_EN.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 35, 50)
MUTED = RGBColor(95, 105, 120)


def style_run(run, bold=False, italic=False, color=None, size=None):
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_sep, fld_text, fld_char_end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for level, size, color in [(1, 16, BLUE), (2, 14, DARK_BLUE), (3, 12, DARK_BLUE)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt({1: 16, 2: 12, 3: 8}[level])
        style.paragraph_format.space_after = Pt({1: 8, 2: 6, 3: 4}[level])
        style.paragraph_format.line_spacing = 1.15


def paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    if first_line:
        p.paragraph_format.first_line_indent = Inches(0.25)
    if text:
        r = p.add_run(text)
        style_run(r, color=INK, size=12)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    style_run(r, bold=True, color=BLUE if level == 1 else DARK_BLUE, size={1: 16, 2: 14, 3: 12}[level])
    return p


def centered(doc, text, size=12, bold=False, color=INK, after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    style_run(r, bold=bold, color=color, size=size)
    return p


def add_title_page(doc):
    centered(doc, "CareerTrackAI", size=24, bold=True, color=BLUE, after=12)
    centered(doc, "A Full-Stack AI-Powered Career Tracking Platform", size=16, bold=True, color=INK, after=24)
    centered(doc, "Project Report", size=14, bold=True, color=DARK_BLUE, after=18)
    centered(doc, "Prepared as a university-style technical report", size=12, color=MUTED, after=40)
    paragraph(
        doc,
        "CareerTrackAI is a web-based platform designed to help students and early-career applicants manage the complete job and internship search process. The system combines application tracking, company and opportunity management, resume analysis, interview preparation, AI-powered guidance, and integrations with external job data providers.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
    )
    centered(doc, "June 2026", size=12, color=MUTED, after=8)
    doc.add_page_break()


def add_contents(doc):
    heading(doc, "Contents", 1)
    for item in [
        "1. Introduction",
        "2. Project Overview",
        "3. Target Users and Objectives",
        "4. System Architecture and Workflow",
        "5. Technologies Used",
        "6. Backend Implementation Overview",
        "7. Frontend Implementation Overview",
        "8. Artificial Intelligence and External APIs",
        "9. Challenges and Solutions",
        "10. Testing and Verification",
        "11. Current Limitations",
        "12. Future Enhancements",
        "13. Conclusion",
    ]:
        paragraph(doc, item, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_contents(doc)

    heading(doc, "1. Introduction", 1)
    paragraph(
        doc,
        "Searching for internships and entry-level jobs is often a fragmented process. Students usually collect company names, job links, application statuses, resume versions, and interview notes across several tools such as spreadsheets, browser bookmarks, email inboxes, and personal notes. This fragmentation makes it difficult to understand progress, identify priorities, prepare for interviews, or learn from previous applications.",
    )
    paragraph(
        doc,
        "CareerTrackAI was developed to solve this problem by providing a unified career operations workspace. The project combines a structured application tracking system with artificial intelligence features that support resume analysis, recommendation generation, interview preparation, link verification, and contextual career guidance. The goal is not only to store information, but also to help the user make better decisions during the job search process.",
    )

    heading(doc, "2. Project Overview", 1)
    paragraph(
        doc,
        "CareerTrackAI is a full-stack web application built around the idea of managing the complete career search pipeline. The user can discover or import opportunities, review company information, save selected rows into a personal workspace, track applications through different stages, upload resumes, analyze resumes using AI, schedule interviews, and receive AI-generated guidance when needed.",
    )
    paragraph(
        doc,
        "The platform is organized around several main modules. The Dashboard gives a high-level view of the user’s progress. Opportunities displays available jobs and internships and allows the user to track a role as an application. Applications uses a Kanban-style board to move each application between Planning, Applied, Interview, Accepted, and Rejected. Resumes manages uploaded CV files and AI analysis results. Interviews stores interview dates, links, locations, and preparation notes. Data Hub manages company and opportunity imports from CSV files and external APIs. AI Studio provides a persistent AI chat experience for general career support.",
    )

    heading(doc, "3. Target Users and Objectives", 1)
    paragraph(
        doc,
        "The primary target users are university students, fresh graduates, and early-career applicants who need a structured way to manage internship and job applications. The system is also useful for career advising environments, where students need guidance in organizing opportunities, preparing resumes, and tracking progress over time.",
    )
    paragraph(
        doc,
        "The main objective of the project is to transform the job search from a scattered manual process into a guided workflow. The system aims to centralize opportunity data, reduce repeated manual tracking, improve resume preparation, provide practical AI assistance, and make each application stage visible and manageable. Another important objective is to keep user data separated by account, so each user has a private career workspace while still being able to optionally save useful shared company records.",
    )

    heading(doc, "4. System Architecture and Workflow", 1)
    paragraph(
        doc,
        "The system follows a client-server architecture. The frontend is built with React and Vite and communicates with the backend using Axios. The backend is built with ASP.NET Core Web API and uses Entity Framework Core to communicate with SQL Server. The backend exposes authenticated API endpoints for users, companies, job opportunities, applications, resumes, interviews, notifications, dashboard statistics, AI features, and usage reporting.",
    )
    paragraph(
        doc,
        "A typical workflow begins when a user registers or logs in. The user can then search for opportunities through Data Hub or import CSV files. Imported data is shown in a review workspace before it is saved. Once an opportunity is saved, it appears in the Opportunities page. The user can track the opportunity, which creates an application record. The application can then be moved through the pipeline. If it reaches the Interview stage, the Interviews page allows the user to schedule the interview and store preparation notes. At any point, the user can use AI Studio, resume analysis, or embedded AI panels for guidance.",
    )

    heading(doc, "5. Technologies Used", 1)
    paragraph(
        doc,
        "The backend uses ASP.NET Core Web API on .NET 10. Entity Framework Core is used for database access and migrations, while SQL Server is used as the relational database. Authentication is implemented using JWT bearer tokens and refresh tokens, and passwords are protected with BCrypt. The backend also uses HttpClient integrations for Gemini, Adzuna, and JobDataLake. Scalar and OpenAPI support are included to make the API easier to inspect and test during development.",
    )
    paragraph(
        doc,
        "The frontend uses React 19 with Vite 7. Tailwind CSS is used for styling, React Router is used for page navigation, Axios is used for API communication, lucide-react provides icons, Recharts provides dashboard and usage charts, dayjs handles date formatting, and react-markdown with remark-gfm is used to render AI-generated Markdown content cleanly. The project also uses oxlint for frontend linting.",
    )

    heading(doc, "6. Backend Implementation Overview", 1)
    paragraph(
        doc,
        "The backend is organized into controllers, services, DTOs, models, and the database context. Controllers define the HTTP endpoints, services contain the main business logic, DTOs define the shape of request and response objects, models represent database entities, and AppDbContext defines the database sets and relationships. This separation makes the backend easier to maintain because each layer has a clear responsibility.",
    )
    paragraph(
        doc,
        "The authentication module handles registration, login, refresh tokens, and logout. The user module stores profile information that can later improve AI recommendations. The company and job opportunity modules manage the user’s career data and support import and export operations. The application module manages the pipeline status for each opportunity. The resume module stores uploaded PDF or DOCX files and connects them to AI analysis. The interview module stores interview title, date, duration, type, meeting link or location, and preparation notes.",
    )
    paragraph(
        doc,
        "The AI service is one of the most important backend components. It communicates with Gemini when the provider is configured and falls back to local guidance when live AI is not available. It supports chat responses, resume analysis, cover letter generation, recommendations, and provider ping checks. The backend also includes usage tracking for Gemini and external data providers, allowing the Usage page to show how different APIs are being used.",
    )

    heading(doc, "7. Frontend Implementation Overview", 1)
    paragraph(
        doc,
        "The frontend was designed as an English-only modern interface. It uses a collapsible sidebar, a sticky header, global search, notification controls, theme settings, and responsive layouts. The design supports both light and dark mode, with additional density preferences that allow the user to choose a compact, comfortable, or spacious layout. The interface was built using Tailwind CSS utilities and reusable components such as cards, buttons, inputs, status pills, and AI panels.",
    )
    paragraph(
        doc,
        "The Dashboard provides a summary of the user’s progress and includes optional charts and AI panels. Applications uses a drag-and-drop Kanban board powered by dnd-kit. Opportunities displays saved opportunities with source badges, job links, tracking actions, delete actions, and AI link verification. Resumes supports file upload, analysis loading states, persistent AI analysis results, and deletion. Interviews supports both scheduled interviews and applications that are only in the Interview stage. AI Studio stores chat history, recommendations, cover letter output, and Gemini token test results. Data Hub organizes imports, external providers, preview rows, editing, deletion, and CSV operations.",
    )
    paragraph(
        doc,
        "A major frontend improvement was persistence. AI chat history, resume analysis results, Data Hub preview rows, and embedded AI panel outputs are stored in localStorage so they do not disappear when the user navigates between pages. This creates a more stable user experience and prevents the loss of generated AI content.",
    )

    heading(doc, "8. Artificial Intelligence and External APIs", 1)
    paragraph(
        doc,
        "Artificial intelligence is integrated as a practical assistant rather than a decorative feature. Gemini is used for career chat, resume analysis, cover letter generation, recommendations, interview preparation notes, job link verification, and AI-powered sourcing guidance. AI output is rendered using Markdown so that bold text, lists, and structured responses are displayed professionally instead of appearing as raw symbols.",
    )
    paragraph(
        doc,
        "External job data is supported through Adzuna and JobDataLake. Adzuna is used to search for opportunities in supported countries, while JobDataLake provides enriched job data through its own API. The system also includes an AI sourcing workflow where the AI creates a search plan and uses the selected provider or a scouting approach related to Google or LinkedIn. Imported data does not automatically enter the user’s workspace without review; it first appears as editable preview rows.",
    )
    paragraph(
        doc,
        "The project also includes usage tracking. The Usage page separates Gemini, Adzuna, and JobDataLake usage and provides a general summary. This is important because AI and external APIs often have quotas, request limits, or future billing implications.",
    )

    heading(doc, "9. Challenges and Solutions", 1)
    paragraph(
        doc,
        "One major challenge was user data separation. At one point, a newly registered account could see companies or opportunities created by another account because some records were treated as shared data. This was solved by adding user scoping through UserId fields and by making shared data optional. A user can now load the shared database manually and save only selected companies into the personal workspace.",
    )
    paragraph(
        doc,
        "Another challenge was dealing with external job APIs. Adzuna does not support every country equally, and some searches may return no results. JobDataLake may also return provider-side errors for some queries. The solution was to expose supported country selection, show clear status messages, preserve preview rows, and allow the user to choose between providers instead of depending on one source.",
    )
    paragraph(
        doc,
        "Resume analysis also required several improvements. A PDF or DOCX file may not always provide readable extracted text. To address this, a resume text extraction service was added, and the analysis logic was improved to avoid empty responses. The frontend now shows loading states, preserves the analysis result, and allows the user to delete old resume files.",
    )
    paragraph(
        doc,
        "The user interface also presented challenges. Dark mode required several color adjustments to make text, cards, buttons, and badges readable. The sidebar needed internal scrolling so that the Settings tab would not disappear on smaller screens. Search results also needed better positioning so they would not overlap with the sidebar. These issues were fixed through layout changes and more consistent Tailwind styling.",
    )
    paragraph(
        doc,
        "AI-generated text initially appeared with raw Markdown symbols such as asterisks. This was solved by using react-markdown and remark-gfm in AI Studio, interview preparation notes, and embedded AI panels. Gemini token errors were also improved so that the user sees understandable messages such as quota reached, invalid API key, or backend connection issue instead of raw technical responses.",
    )

    heading(doc, "10. Testing and Verification", 1)
    paragraph(
        doc,
        "The project was verified through both build-time checks and functional user-flow testing. The frontend was checked with npm run lint and npm run build. The backend was built using dotnet build, sometimes with a temporary output directory because the running API process locked the default executable file. These checks confirmed that the frontend compiled successfully and that backend code changes were valid.",
    )
    paragraph(
        doc,
        "Functional testing focused on realistic user scenarios. These included registering a new account, confirming that user-specific data does not leak from another account, importing opportunities, tracking an opportunity as an application, moving applications between statuses, ensuring Interview-stage applications appear on the Interviews page, scheduling interview details, analyzing resumes, and testing Gemini token status. The system was also reviewed for dark mode readability and persistence of AI-generated results.",
    )

    heading(doc, "11. Current Limitations", 1)
    paragraph(
        doc,
        "The payment and subscription flow currently exists as a prepared structure rather than a complete production feature. The application includes plan screens and checkout draft behavior, but it is not yet connected to a real payment gateway such as Stripe, Tap Payments, or Moyasar. This was intentionally postponed because payment integration requires account setup, gateway configuration, backend webhook handling, and security review.",
    )
    paragraph(
        doc,
        "Another limitation is the temporary handling of personal API keys on the frontend. A production system should store user-provided API keys in an encrypted backend vault rather than localStorage. The system also depends on external provider quotas, especially Gemini daily limits, which means some AI features may become unavailable until the quota resets or another key is configured.",
    )
    paragraph(
        doc,
        "The frontend build currently produces a large JavaScript bundle warning. This does not prevent the project from running, but it suggests that future optimization should include code splitting and route-based lazy loading. External job APIs may also return limited results depending on the country, provider support, and query wording.",
    )

    heading(doc, "12. Future Enhancements", 1)
    paragraph(
        doc,
        "A key future enhancement is a real subscription and payment system. This would include connecting a payment provider, confirming payments through backend webhooks, updating the user plan only after confirmed payment, and linking AI quotas to the selected plan. The planned packages could include a free plan, a managed AI plan, and a Bring Your Own Key plan.",
    )
    paragraph(
        doc,
        "Another future improvement is secure API key management. Instead of storing user keys in the browser, the backend should encrypt and store them securely. The system could then support per-user Gemini keys, provider switching, and usage limits. This would make the platform safer and more suitable for production deployment.",
    )
    paragraph(
        doc,
        "Additional features could include email and calendar reminders, admin dashboards, advanced analytics, role-based access for university career centers, AI-powered resume versioning, better matching between resumes and job descriptions, and deployment through a CI/CD pipeline. A vector database such as Pinecone or another retrieval system could also be used in the future to make the AI more aware of user documents and historical application data.",
    )

    heading(doc, "13. Conclusion", 1)
    paragraph(
        doc,
        "CareerTrackAI is a comprehensive full-stack project that addresses a real problem faced by students and early-career applicants. It combines structured career tracking with useful artificial intelligence features and external job data integrations. The system provides a practical workflow from opportunity discovery to application tracking, resume analysis, and interview preparation.",
    )
    paragraph(
        doc,
        "The project demonstrates the use of modern frontend development, ASP.NET Core backend architecture, database design, authentication, AI integration, external APIs, and user-centered interface improvements. Although payment integration and production-grade key storage remain future work, the current system forms a strong foundation for a scalable AI-powered career management platform.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_document()
    print(OUT)
